"""Generate PROVISO_INTRO.docx — styled Word document for management presentation."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colors ─────────────────────────────────────────────────────
NAVY      = RGBColor(0x03, 0x09, 0x10)
DARK_BLUE = RGBColor(0x0A, 0x18, 0x28)
ACCENT    = RGBColor(0x15, 0x65, 0xD8)
LIGHT_BLU = RGBColor(0x4A, 0x9F, 0xFF)
GREEN     = RGBColor(0x00, 0xC8, 0x70)
GOLD      = RGBColor(0xF0, 0xA5, 0x00)
RED       = RGBColor(0xFF, 0x3D, 0x5A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
MID_GRAY  = RGBColor(0x58, 0x78, 0xA0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x2E)
SUBTLE    = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "CCCCCC"))
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def styled_table(doc, headers, rows, header_bg="0A1828", accent_col=None):
    """Create a styled table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.upper())
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_BLU
        run.font.bold = True
        run.font.name = "Consolas"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_bg)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

            if accent_col is not None and ci == accent_col:
                run.font.color.rgb = GREEN
                run.font.bold = True
            elif ci == 0:
                run.font.color.rgb = SUBTLE
            else:
                run.font.color.rgb = TEXT_DARK

            if ri % 2 == 0:
                set_cell_shading(cell, "F5F8FC")
            else:
                set_cell_shading(cell, "FFFFFF")

    # Borders
    for row in table.rows:
        for cell in row.cells:
            border = {"bottom": {"val": "single", "sz": "4", "color": "DEE5EF"}}
            set_cell_border(cell, **border)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a colored heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = ACCENT
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = ACCENT
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = LIGHT_BLU
            run.font.size = Pt(13)
        run.font.name = "Calibri"
    return h


def add_body(doc, text, bold_phrases=None, color=None):
    """Add a body paragraph with optional bold phrases."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)

    if bold_phrases:
        parts = text.split("**")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"
            run.font.color.rgb = color or TEXT_DARK
            if i % 2 == 1:
                run.font.bold = True
                run.font.color.rgb = ACCENT
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = color or TEXT_DARK
    return p


def add_advantage_block(doc, icon, title, desc):
    """Add an advantage as a styled block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(f"{icon}  {title}")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Inches(0.3)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = SUBTLE
    run2.font.name = "Calibri"


def add_step_block(doc, num, title, desc):
    """Add a how-it-works step."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    run_num = p.add_run(f"  {num}  ")
    run_num.font.size = Pt(11)
    run_num.font.bold = True
    run_num.font.color.rgb = WHITE
    run_num.font.highlight_color = 2  # Not ideal but close

    run_title = p.add_run(f"   {title}")
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_title.font.color.rgb = ACCENT
    run_title.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Inches(0.3)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = SUBTLE
    run2.font.name = "Calibri"


# ══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── TITLE ──────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
run = title_p.add_run("proviso")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = ACCENT
run.font.name = "Calibri"

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
run = sub_p.add_run("WORKFLOW INGESTION FOR M-FILES")
run.font.size = Pt(11)
run.font.color.rgb = MID_GRAY
run.font.name = "Consolas"

phase_p = doc.add_paragraph()
phase_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
phase_p.paragraph_format.space_after = Pt(6)
run = phase_p.add_run("Phase I  ·  Proof of Concept")
run.font.size = Pt(10)
run.font.color.rgb = LIGHT_BLU
run.font.name = "Calibri"
run.font.italic = True

link_p = doc.add_paragraph()
link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
link_p.paragraph_format.space_after = Pt(4)
run = link_p.add_run("▸ Live Demo:  ")
run.font.size = Pt(10)
run.font.color.rgb = SUBTLE
run.font.name = "Calibri"
run2 = link_p.add_run("provisio-theta.vercel.app")
run2.font.size = Pt(10)
run2.font.color.rgb = GREEN
run2.font.bold = True
run2.font.name = "Consolas"

# Divider
div_p = doc.add_paragraph()
div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = div_p.add_run("━" * 60)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xDE, 0xE5, 0xEF)

# ── THE PROBLEM ────────────────────────────────────────────────
add_heading_styled(doc, "The Problem We're Solving", 1)

add_body(doc, "Today, every M-Files workflow deployment requires **triple entry** — the same logic written three times, in three different places, by hand:", bold_phrases=True)

styled_table(doc,
    ["Step", "Manual Task", "Risk"],
    [
        ["1", "Write the SOW document", "⚠ Misinterpretation"],
        ["2", "Draw a workflow diagram", "⚠ Drift from SOW"],
        ["3", "Rebuild it in M-Files Admin", "⚠ Fat-finger errors"],
    ]
)

doc.add_paragraph()
add_body(doc, "Each step is disconnected. A change in one doesn't update the others. Mistakes compound. Consultants spend hours clicking through admin screens reproducing what's already on paper.")
add_body(doc, "**Proviso eliminates all three steps and replaces them with one.**", bold_phrases=True)

# ── THE SOLUTION ───────────────────────────────────────────────
add_heading_styled(doc, "The Solution: SOW-to-Vault Automation", 1)

add_body(doc, "The SOW spreadsheet becomes the **single source of truth**. Everything else — the diagram, the PRD, the vault configuration — is generated automatically from that one input.", bold_phrases=True)

# Callout box
callout = doc.add_paragraph()
callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
callout.paragraph_format.space_before = Pt(14)
callout.paragraph_format.space_after = Pt(14)
r1 = callout.add_run("One input.  ")
r1.font.size = Pt(14); r1.font.bold = True; r1.font.color.rgb = ACCENT; r1.font.name = "Calibri"
r2 = callout.add_run("Four outputs.  ")
r2.font.size = Pt(14); r2.font.bold = True; r2.font.color.rgb = TEXT_DARK; r2.font.name = "Calibri"
r3 = callout.add_run("Zero re-entry.")
r3.font.size = Pt(14); r3.font.bold = True; r3.font.color.rgb = GREEN; r3.font.name = "Calibri"

# Data Flow
add_heading_styled(doc, "Data Flow", 3)

flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow.paragraph_format.space_before = Pt(6)
flow.paragraph_format.space_after = Pt(6)
lines = [
    "Spreadsheet (React UI)",
    "│",
    "▼",
    "workflow.json  ◄── the single source of truth",
    "│",
    "┌────────┼────────────────┐",
    "▼              ▼                        ▼",
    "Diagram      PRD (.md)        M-Files Vault",
    "(Mermaid)    (NLP output)     (COM API)",
]
for line in lines:
    run = flow.add_run(line + "\n")
    run.font.size = Pt(9.5)
    run.font.name = "Consolas"
    if "workflow.json" in line:
        run.font.color.rgb = ACCENT
        run.font.bold = True
    elif "single source" in line:
        run.font.color.rgb = LIGHT_BLU
    else:
        run.font.color.rgb = MID_GRAY

# How it works
add_heading_styled(doc, "How It Works", 2)

add_step_block(doc, "1", "SOW Editor",
    "Define states, transitions, users, and properties in a structured spreadsheet. "
    "A live Mermaid diagram renders in real time as you type. If the logic looks wrong "
    "in the picture, you fix it in the spreadsheet — immediately.")

add_step_block(doc, "2", "Generate PRD",
    "Local NLP (regex + pattern matching) transforms the technical spreadsheet data "
    "into a client-ready Product Requirements Document. An AI-enhanced option via "
    "Claude is available for more complex narrative requirements.")

add_step_block(doc, "3", "Ingest Workflow",
    "The COM API takes the workflow JSON and writes the structure directly into the "
    "M-Files Vault — states, transitions, and aliases — in seconds. The consultant "
    "then opens M-Files Admin and adds only the business rules and conditions.")

# ── ADVANTAGES ─────────────────────────────────────────────────
add_heading_styled(doc, "Why This Is a Game-Changer", 1)

add_advantage_block(doc, "🎯", "Zero Redundancy",
    "Build the SOW once. The diagram, PRD, JSON, and vault configuration are all "
    "byproducts of that single effort. No remapping. No re-entry. No copy-paste.")

add_advantage_block(doc, "🔒", "Zero Drift",
    "The diagram and vault are generated from the same source. They cannot go out "
    "of sync. What the client approved is exactly what gets built.")

add_advantage_block(doc, "✅", "Zero Fat-Finger Errors",
    "Manual data entry into M-Files Admin is replaced by automated ingestion. "
    "11 states in the SOW = exactly 11 states in the vault. No typos, no missed connections.")

add_advantage_block(doc, "⚡", "Instant Iteration",
    "Customer changes their mind about a 'Review' state? Change one cell in the "
    "spreadsheet. The diagram, PRD, and vault configuration all update. What took "
    "an hour now takes seconds.")

add_advantage_block(doc, "🤝", "Client Transparency",
    "Clients see a clean, professional workflow diagram early in the engagement — "
    "tied directly to the SOW they signed. No gap between promise and delivery.")

add_advantage_block(doc, "🚀", "Consultant Productivity",
    "Stop spending time on mechanical data entry. Start spending time on what matters: "
    "business rules, permissions, and client-specific logic. Proviso handles the "
    "scaffolding — the consultant handles the intelligence.")

# ── TIME SAVINGS ───────────────────────────────────────────────
add_heading_styled(doc, "Time Savings Estimate", 1)

styled_table(doc,
    ["Task", "Before (Manual)", "After (Proviso)", "Saved"],
    [
        ["Define workflow states + transitions", "45–60 min", "5 min", "~50 min"],
        ["Draw workflow diagram", "30–45 min", "0 min (auto)", "~35 min"],
        ["Write PRD documentation", "60–90 min", "1 min (NLP)", "~75 min"],
        ["Configure M-Files Admin", "60–120 min", "30 sec (COM API)", "~90 min"],
        ["Rework after client changes", "30–60 min/change", "2 min/change", "~45 min"],
        ["TOTAL PER WORKFLOW", "4–6 hours", "~10 minutes", "~5 HOURS"],
    ],
    accent_col=3
)

doc.add_paragraph()
add_body(doc, "For a typical project with 2–3 workflows, that's **10–15 hours saved per engagement.**", bold_phrases=True)

# ── PHASE I SCOPE ──────────────────────────────────────────────
add_heading_styled(doc, "Phase I Scope", 1)

add_body(doc, "This proof of concept focuses on the **backbone** — ensuring that a state defined in the editor successfully appears as a state in the vault.", bold_phrases=True)

styled_table(doc,
    ["In Scope (Phase I)", "Out of Scope (Phase II)"],
    [
        ["✅  Workflow creation", "Phase II — Business rules / conditions"],
        ["✅  State definitions", "Phase II — Automatic state transitions"],
        ["✅  Transition mapping", "Phase II — Permission assignments"],
        ["✅  Alias assignment", "Phase II — Script-based triggers"],
        ["✅  Live diagram preview", "Phase II — Multi-workflow orchestration"],
        ["✅  PRD generation (NLP + AI)", "Phase II — Class / object type creation"],
    ]
)

# ── TECHNOLOGY ─────────────────────────────────────────────────
add_heading_styled(doc, "Technology", 1)

styled_table(doc,
    ["Component", "Stack"],
    [
        ["Frontend", "React 18 · Vite · Mermaid.js"],
        ["Data Format", "JSON — spreadsheet converts to workflow.json, driving diagram, PRD, and vault"],
        ["PRD Output", "Markdown (.md) — generated from JSON by NLP engine"],
        ["PRD Engine", "Local NLP (regex + pattern matching) · Claude AI (optional)"],
        ["Vault Ingestion", "Python · Flask · pywin32 · M-Files COM API"],
        ["Hosting (demo)", "Vercel"],
    ]
)

# ── FOOTER ─────────────────────────────────────────────────────
div_p2 = doc.add_paragraph()
div_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
div_p2.paragraph_format.space_before = Pt(20)
run = div_p2.add_run("━" * 60)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xDE, 0xE5, 0xEF)

demo_p = doc.add_paragraph()
demo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
demo_p.paragraph_format.space_after = Pt(4)
run = demo_p.add_run("▸ Live Demo → ")
run.font.size = Pt(11); run.font.color.rgb = SUBTLE; run.font.name = "Calibri"
run2 = demo_p.add_run("provisio-theta.vercel.app")
run2.font.size = Pt(11); run2.font.color.rgb = GREEN; run2.font.bold = True; run2.font.name = "Consolas"

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author_p.add_run("Author: Harry Joseph  ·  Phase I POC  ·  May 2026")
run.font.size = Pt(9)
run.font.color.rgb = MID_GRAY
run.font.name = "Consolas"

# ── SAVE ───────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "PROVISO_INTRO.docx")
doc.save(output_path)
print(f"OK - Created: {output_path}")
